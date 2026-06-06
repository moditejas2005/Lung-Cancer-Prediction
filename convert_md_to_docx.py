import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_formatted_text(paragraph, text, italic=False, is_code=False):
    # Matches markdown bold (**bold**) and inline code (`code`)
    pattern = r'(\*\*.*?\*\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            run = paragraph.add_run(part)
            
        if not (part.startswith('`') and part.endswith('`')):
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            
        if italic:
            run.font.italic = True

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_text = ""
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # ─── Code Block Handler ───
        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                run = p.add_run(code_text.rstrip())
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F1F5F9')
                pPr.append(shd)
                
                code_text = ""
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text += line
            i += 1
            continue
            
        # ─── Table Handler ───
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            in_table = False
            parsed_rows = []
            for row in table_rows:
                cells = [c.strip() for c in row.split("|")[1:-1]]
                parsed_rows.append(cells)
            
            # Remove separator rows like |---|---|
            filtered_rows = []
            for row in parsed_rows:
                if len(row) > 0 and all(all(char == '-' for char in col) for col in row if col):
                    continue
                filtered_rows.append(row)
                
            if filtered_rows:
                col_count = max(len(r) for r in filtered_rows)
                t = doc.add_table(rows=len(filtered_rows), cols=col_count)
                t.style = 'Table Grid'
                for r_idx, row_data in enumerate(filtered_rows):
                    row = t.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.text = cell_value
                            if r_idx == 0:
                                set_cell_background(cell, "1F4E79")
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.bold = True
                                        r.font.color.rgb = RGBColor(255, 255, 255)
            table_rows = []
            
        if not stripped:
            i += 1
            continue
            
        # ─── Headings ───
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.font.color.rgb = RGBColor(31, 78, 121)
                r.font.name = 'Calibri'
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.color.rgb = RGBColor(46, 117, 182)
                r.font.name = 'Calibri'
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
                r.font.name = 'Calibri'
        elif stripped.startswith("#### "):
            p = doc.add_heading(stripped[5:], level=4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = 'Calibri'
                
        # ─── Bullet Lists ───
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            add_formatted_text(p, stripped[2:])
            
        # ─── Blockquotes ───
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, stripped[2:], italic=True)
            
        # ─── Numbered Lists ───
        elif re.match(r"^\d+\.\s*", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            num_part = re.match(r"^\d+\.\s*", stripped).group(0)
            text_part = stripped[len(num_part):]
            p.add_run(num_part).font.bold = True
            add_formatted_text(p, text_part)
            
        # ─── Horizontal Rules ───
        elif stripped.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("____________________________________________________")
            run.font.color.rgb = RGBColor(192, 192, 192)
            
        # ─── Standard Paragraph ───
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, stripped)
            
        i += 1
        
    doc.save(docx_path)
    print(f"Successfully saved {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx("local_demo_guide.md", "local_demo_guide.docx")
    convert_md_to_docx("demo_presentation_speech.md", "demo_presentation_speech.docx")
